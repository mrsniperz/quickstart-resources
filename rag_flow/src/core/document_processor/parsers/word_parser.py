"""
模块名称: word_parser
功能描述: Word文档解析器，基于python-docx实现Word文档的文本提取、表格识别和格式保持
创建日期: 2024-01-15
作者: Sniperz
版本: v1.0.0
"""

import logging
from typing import Dict, List, Optional, Any
from pathlib import Path
from dataclasses import dataclass
from docx import Document
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph

from ..extractors.metadata_extractor import MetadataExtractor
from ..extractors.table_extractor import TableExtractor


@dataclass
class WordParseResult:
    """Word解析结果数据类"""
    text_content: str
    metadata: Dict[str, Any]
    tables: List[Dict[str, Any]]
    paragraphs: List[Dict[str, Any]]
    structure_info: Dict[str, Any]


class WordParser:
    """
    Word文档处理器
    
    基于python-docx实现Word文档的全面解析，包括：
    - 文本内容提取
    - 表格数据提取
    - 段落结构分析
    - 格式保持与转换
    - 元数据提取
    """
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        初始化Word解析器
        
        Args:
            config (dict, optional): 配置参数
                - preserve_formatting (bool): 是否保持格式，默认True
                - extract_tables (bool): 是否提取表格，默认True
                - extract_headers_footers (bool): 是否提取页眉页脚，默认False
        """
        self.config = config or {}
        self.logger = logging.getLogger(__name__)
        
        # 初始化提取器
        self.metadata_extractor = MetadataExtractor()
        self.table_extractor = TableExtractor()
        
        # 配置参数
        self.preserve_formatting = self.config.get('preserve_formatting', True)
        self.extract_tables = self.config.get('extract_tables', True)
        self.extract_headers_footers = self.config.get('extract_headers_footers', False)
        
    def parse(self, file_path: str) -> WordParseResult:
        """
        解析Word文档
        
        Args:
            file_path (str): Word文件路径
            
        Returns:
            WordParseResult: 解析结果
            
        Raises:
            FileNotFoundError: 文件不存在
            ValueError: 文件格式不支持
            Exception: 解析过程中的其他错误
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"Word文件不存在: {file_path}")
                
            if file_path.suffix.lower() not in ['.docx', '.doc']:
                raise ValueError(f"不支持的文件格式: {file_path.suffix}")
                
            self.logger.info(f"开始解析Word文档: {file_path}")
            
            # 打开文档
            doc = Document(str(file_path))
            
            # 提取元数据
            metadata = self._extract_metadata(doc, file_path)
            
            # 提取段落内容
            paragraphs = self._extract_paragraphs(doc)
            
            # 提取文本内容
            text_content = self._extract_text(paragraphs)
            
            # 提取表格
            tables = []
            if self.extract_tables:
                tables = self._extract_tables(doc)
            
            # 分析文档结构
            structure_info = self._analyze_structure(doc, paragraphs)
            
            result = WordParseResult(
                text_content=text_content,
                metadata=metadata,
                tables=tables,
                paragraphs=paragraphs,
                structure_info=structure_info
            )
            
            self.logger.info(f"Word解析完成: {len(paragraphs)}个段落, {len(tables)}个表格")
            return result
            
        except Exception as e:
            self.logger.error(f"Word解析失败: {e}")
            raise
    
    def _extract_metadata(self, doc: Document, file_path: Path) -> Dict[str, Any]:
        """
        提取文档元数据
        
        Args:
            doc: python-docx文档对象
            file_path: 文件路径
            
        Returns:
            dict: 元数据信息
        """
        try:
            core_props = doc.core_properties
            
            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'comments': core_props.comments or '',
                'category': core_props.category or '',
                'created': core_props.created,
                'modified': core_props.modified,
                'last_modified_by': core_props.last_modified_by or '',
                'revision': core_props.revision,
                'version': core_props.version or '',
                'file_path': str(file_path),
                'file_size': file_path.stat().st_size if file_path.exists() else 0,
                'file_extension': file_path.suffix.lower(),
                'is_word_document': True
            }
            
            return metadata
            
        except Exception as e:
            self.logger.error(f"元数据提取失败: {e}")
            return {'file_path': str(file_path), 'is_word_document': True}
    
    def _extract_paragraphs(self, doc: Document) -> List[Dict[str, Any]]:
        """
        提取文档段落
        
        Args:
            doc: python-docx文档对象
            
        Returns:
            list: 段落信息列表
        """
        try:
            paragraphs = []
            
            for para_idx, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():  # 跳过空段落
                    para_info = {
                        'index': para_idx,
                        'text': paragraph.text,
                        'style': paragraph.style.name if paragraph.style else 'Normal',
                        'alignment': str(paragraph.alignment) if paragraph.alignment else None,
                        'runs': []
                    }
                    
                    # 提取运行信息（格式化文本片段）
                    for run in paragraph.runs:
                        if run.text.strip():
                            run_info = {
                                'text': run.text,
                                'bold': run.bold,
                                'italic': run.italic,
                                'underline': run.underline,
                                'font_name': run.font.name,
                                'font_size': run.font.size.pt if run.font.size else None
                            }
                            para_info['runs'].append(run_info)
                    
                    paragraphs.append(para_info)
            
            return paragraphs
            
        except Exception as e:
            self.logger.error(f"段落提取失败: {e}")
            return []
    
    def _extract_text(self, paragraphs: List[Dict[str, Any]]) -> str:
        """
        从段落中提取纯文本内容
        
        Args:
            paragraphs: 段落信息列表
            
        Returns:
            str: 提取的文本内容
        """
        try:
            text_parts = []
            
            for para in paragraphs:
                if self.preserve_formatting:
                    # 保持段落结构
                    style = para.get('style', 'Normal')
                    text = para['text']
                    
                    if style.startswith('Heading'):
                        # 标题添加标记
                        level = style.replace('Heading ', '') if 'Heading ' in style else '1'
                        text_parts.append(f"{'#' * int(level) if level.isdigit() else '#'} {text}")
                    else:
                        text_parts.append(text)
                else:
                    # 仅提取纯文本
                    text_parts.append(para['text'])
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            self.logger.error(f"文本提取失败: {e}")
            return ""
    
    def _extract_tables(self, doc: Document) -> List[Dict[str, Any]]:
        """
        提取文档中的表格
        
        Args:
            doc: python-docx文档对象
            
        Returns:
            list: 表格数据列表
        """
        try:
            all_tables = []
            
            for table_idx, table in enumerate(doc.tables):
                try:
                    # 提取表格数据
                    table_data = []
                    
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            row_data.append(cell.text.strip())
                        table_data.append(row_data)
                    
                    if table_data:
                        table_info = {
                            'table_index': table_idx,
                            'data': table_data,
                            'rows': len(table_data),
                            'columns': len(table_data[0]) if table_data else 0,
                            'style': table.style.name if table.style else None
                        }
                        all_tables.append(table_info)
                        
                except Exception as e:
                    self.logger.warning(f"表格{table_idx}提取失败: {e}")
                    continue
            
            return all_tables
            
        except Exception as e:
            self.logger.error(f"表格提取失败: {e}")
            return []
    
    def _analyze_structure(self, doc: Document, paragraphs: List[Dict[str, Any]]) -> Dict[str, Any]:
        """
        分析文档结构
        
        Args:
            doc: python-docx文档对象
            paragraphs: 段落信息列表
            
        Returns:
            dict: 结构分析信息
        """
        try:
            structure = {
                'total_paragraphs': len(paragraphs),
                'total_tables': len(doc.tables),
                'heading_structure': [],
                'styles_used': set(),
                'word_count': 0,
                'character_count': 0
            }
            
            # 分析标题结构
            for para in paragraphs:
                style = para.get('style', 'Normal')
                structure['styles_used'].add(style)
                
                if style.startswith('Heading'):
                    structure['heading_structure'].append({
                        'level': style,
                        'text': para['text'],
                        'index': para['index']
                    })
                
                # 统计字数
                text = para['text']
                structure['word_count'] += len(text.split())
                structure['character_count'] += len(text)
            
            structure['styles_used'] = list(structure['styles_used'])
            
            return structure
            
        except Exception as e:
            self.logger.error(f"结构分析失败: {e}")
            return {}
    
    def extract_paragraph_text(self, file_path: str, paragraph_index: int) -> str:
        """
        提取指定段落的文本
        
        Args:
            file_path (str): Word文件路径
            paragraph_index (int): 段落索引
            
        Returns:
            str: 段落文本内容
            
        Raises:
            IndexError: 段落索引超出范围
        """
        try:
            doc = Document(file_path)
            
            if paragraph_index >= len(doc.paragraphs):
                raise IndexError(f"段落索引超出范围: {paragraph_index} >= {len(doc.paragraphs)}")
            
            return doc.paragraphs[paragraph_index].text
            
        except Exception as e:
            self.logger.error(f"段落文本提取失败: {e}")
            raise
    
    def get_supported_formats(self) -> List[str]:
        """
        获取支持的文件格式
        
        Returns:
            list: 支持的文件扩展名列表
        """
        return ['.docx', '.doc']
